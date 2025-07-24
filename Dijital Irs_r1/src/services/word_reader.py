"""
Belirtilen word dosyasından tablo okuma işlemlerini yapar
"""
import pandas as pd
from docx import Document
from typing import List

class WordReaderService:
    def __init__(self):
        self.current_document = None

    def load_document(self, file_path: str) -> bool:
        """Word Dosyasını Yükler"""
        try:
            doc = Document(file_path)
            self.current_document = doc
            print(f"✓ Word dosyası yüklendi: {file_path}")
            return True
        except Exception as e:
            print(f"HATA: Word dosyası yüklenemedi - {e}")
            return False

    def extract_tables(self, file_path: str) -> List:
        """
        Word içerisindeki Tabloları Veri Toplar ve Liste Olarak Döndürür
        """
        print("Veri işleme başlıyor...")
        
        # Header'ı ilk eleman olarak liste içinde tanımla
        headers = ["ITEM NO", "DIMENSION", "ACTUAL", "BADGE", "TOOLING", "REMARKS", "B/P ZONE", "INSP. LEVEL"]
        extracted_data = [headers]  # İlk eleman header LİSTESİ
        
        try:
            # Word dokümanını yükle
            word_document = self.load_document(file_path=file_path)
            if not word_document:
                return []
            
            tables = self.current_document.tables
            print(f"  {len(tables)} tablo bulundu")
            
            for table_idx, table in enumerate(tables):
                print(f"  Tablo {table_idx + 1} işleniyor...")
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Hücredeki tüm paragrafları birleştir
                        cell_text = '\n'.join([p.text for p in cell.paragraphs])
                        row_data.append(cell_text)
                    
                    # DEBUG: Satır verilerini kontrol et
                    print(f"    Debug - row_data uzunluğu: {len(row_data)}")
                    if len(row_data) > 0:
                        print(f"    Debug - İlk element: '{row_data[0]}'")
                    
                    # Filtreleme mantığı
                    if (len(row_data) > 2 and 
                        row_data[0].startswith("KN") and 
                        not row_data[1].strip().endswith("Inch")):
                        
                        # Satırı 8 kolona pad et veya kırp
                        padded_row = row_data + [''] * (len(headers) - len(row_data))
                        padded_row = padded_row[:len(headers)]
                        
                        extracted_data.append(padded_row)  # LİSTE ekleniyor
                        print(f"    ✓ Satır eklendi: {row_data[0]} (kolon sayısı: {len(padded_row)})")
            
            print(f"✓ Toplam {len(extracted_data) - 1} karakter çıkarıldı")
            print(f"Header: {extracted_data[0]}")
            if len(extracted_data) > 1:
                print(f"İlk veri satırı: {extracted_data[1]} (uzunluk: {len(extracted_data[1])})")
            
            return extracted_data
            
        except Exception as e:
            print(f"HATA: Tablo çıkarma işleminde sorun: {e}")
            return []