"""
Word Save As Servisi - Actual değerlerini Word tablosuna yazıp kaydetme
services/word_save_as.py
"""
import os
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from typing import List, Tuple, Optional
from tkinter import filedialog, messagebox
from .data_processor import TeknikResimKarakteri

class WordSaveAsService:
    """
    Word dosyasını actual değerleriyle birlikte yeni bir lokasyona kaydetme servisi
    Tolerans kontrollü kırmızı bold formatı ile
    """
    
    def __init__(self):
        self.current_document = None
        self.original_file_path = None
        
    def load_original_document(self, file_path: str) -> bool:
        """Orijinal Word dosyasını yükler"""
        try:
            self.current_document = Document(file_path)
            self.original_file_path = file_path
            print(f"✓ Orijinal Word dosyası yüklendi: {file_path}")
            return True
        except Exception as e:
            print(f"HATA: Word dosyası yüklenemedi - {e}")
            return False
    
    def find_table_columns(self, table) -> tuple:
        """Tabloda ITEM NO ve ACTUAL kolonlarının indekslerini bulur"""
        actual_col_index = None
        item_no_col_index = None
        
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for col_idx, cell in enumerate(header_row.cells):
                cell_text = cell.text.strip().upper()
                if "ACTUAL" in cell_text:
                    actual_col_index = col_idx
                if "ITEM" in cell_text and "NO" in cell_text:
                    item_no_col_index = col_idx
        
        return item_no_col_index, actual_col_index
    
    def parse_multiple_values(self, actual_value: str) -> List[str]:
        """
        Birden fazla değeri parse eder (örn: '25.4/25.6' -> ['25.4', '25.6'])
        """
        if not actual_value:
            return []
        
        # "/" ile ayrılmış değerleri ayır
        values = [v.strip() for v in actual_value.split('/')]
        return [v for v in values if v]  # Boş stringleri filtrele
    
    def check_tolerance(self, value_str: str, karakter: TeknikResimKarakteri) -> Tuple[bool, str]:
        """
        Tek bir değerin tolerans durumunu kontrol eder
        
        Returns:
            Tuple[bool, str]: (tolerans_içinde_mi, durum_mesajı)
        """
        try:
            # Sayısal değere çevir
            value = float(value_str.replace(',', '.'))
            
            # Tolerans bilgileri var mı?
            has_lower = hasattr(karakter, 'lower_limit') and karakter.lower_limit is not None
            has_upper = hasattr(karakter, 'upper_limit') and karakter.upper_limit is not None
            
            if not has_lower and not has_upper:
                return True, "Tolerans tanımlanmamış"
            
            # Tolerans kontrolü
            if has_lower and has_upper:
                in_tolerance = karakter.lower_limit <= value <= karakter.upper_limit
                status = f"Limit: {karakter.lower_limit} ↔ {karakter.upper_limit}"
            elif has_upper:
                in_tolerance = value <= karakter.upper_limit
                status = f"Max: {karakter.upper_limit}"
            elif has_lower:
                in_tolerance = value >= karakter.lower_limit
                status = f"Min: {karakter.lower_limit}"
            else:
                return True, "Limit yok"
            
            return in_tolerance, status
            
        except ValueError:
            # Sayısal olmayan değer
            return True, "Sayısal olmayan değer"
    
    def check_multiple_values_tolerance(self, actual_value: str, karakter: TeknikResimKarakteri) -> Tuple[List[bool], str]:
        """
        Birden fazla değerin tolerans durumunu kontrol eder
        
        Returns:
            Tuple[List[bool], str]: (her_değer_için_tolerans_durumu, genel_durum_mesajı)
        """
        values = self.parse_multiple_values(actual_value)
        
        if not values:
            return [], "Değer yok"
        
        tolerance_results = []
        status_messages = []
        
        for i, value in enumerate(values):
            in_tolerance, status = self.check_tolerance(value, karakter)
            tolerance_results.append(in_tolerance)
            
            value_status = "✅" if in_tolerance else "❌"
            status_messages.append(f"Değer {i+1} ({value}): {value_status}")
        
        # Genel durum mesajı
        any_out_of_tolerance = not all(tolerance_results)
        general_status = " | ".join(status_messages)
        
        return tolerance_results, general_status
    
    def apply_yellow_highlight(self, cell, text: str, highlight_parts: List[bool] = None):
        """
        Word hücresine sarı highlight uygular (Stack Overflow yöntemi)
        
        Args:
            cell: Word tablosu hücresi
            text: Yazılacak metin
            highlight_parts: Hangi bölümlerin işaretleneceği (/ ile ayrılmış değerler için)
        """
        try:
            # Hücreyi temizle
            cell.text = ""
            paragraph = cell.paragraphs[0]
            
            if not highlight_parts:
                # Tüm metni sarı highlight yap
                run = paragraph.add_run(text)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                print(f"      🟡 Tüm metin sarı highlight yapıldı: {text}")
                            
            else:
                # Birden fazla değer için kısmi highlighting
                values = self.parse_multiple_values(text)
                
                for i, (value, should_highlight) in enumerate(zip(values, highlight_parts)):
                    if i > 0:
                        # Ayırıcı ekle (normal format)
                        paragraph.add_run(" / ")
                    
                    # Değeri ekle
                    run = paragraph.add_run(value)
                    
                    if should_highlight:
                        # Sarı highlight
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        print(f"      🟡 Değer {i+1} sarı highlight yapıldı: {value}")
                    else:
                        # Normal format
                        print(f"      ✅ Değer {i+1} normal: {value}")
                        
        except Exception as e:
            print(f"      ⚠ Sarı highlight hatası: {e}")
            # Fallback: Sadece metni yaz
            cell.text = text
            print(f"      🟡 [FALLBACK] Sarı highlight olması gereken metin yazıldı: {text}")
    
    def apply_red_bold_format(self, cell, text: str, highlight_parts: List[bool] = None):
        """
        Word hücresine kırmızı ve bold format uygular
        
        Args:
            cell: Word tablosu hücresi
            text: Yazılacak metin
            highlight_parts: Hangi bölümlerin formatlanacağı (/ ile ayrılmış değerler için)
        """
        try:
            # Hücreyi temizle
            cell.text = ""
            paragraph = cell.paragraphs[0]
            
            if not highlight_parts:
                # Tüm metni kırmızı bold yap
                run = paragraph.add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # Kırmızı
                print(f"      🔴 Tüm metin kırmızı bold yapıldı: {text}")
                            
            else:
                # Birden fazla değer için kısmi formatlama
                values = self.parse_multiple_values(text)
                
                for i, (value, should_format) in enumerate(zip(values, highlight_parts)):
                    if i > 0:
                        # Ayırıcı ekle (normal format)
                        paragraph.add_run(" / ")
                    
                    # Değeri ekle
                    run = paragraph.add_run(value)
                    
                    if should_format:
                        # Kırmızı bold format
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Kırmızı
                        print(f"      🔴 Değer {i+1} kırmızı bold yapıldı: {value}")
                    else:
                        # Normal format
                        print(f"      ✅ Değer {i+1} normal: {value}")
                        
        except Exception as e:
            print(f"      ⚠ Kırmızı bold format hatası: {e}")
            # Fallback: Sadece metni yaz
            cell.text = text
            print(f"      🔴 [FALLBACK] Kırmızı bold olması gereken metin yazıldı: {text}")
    
    def apply_yellow_highlight_fallback(self, cell, text: str):
        """
        Sarı highlight alternatifi - hücre arka planını sarı yapar
        """
        try:
            # Metni yaz
            cell.text = text
            
            # Hücre arka planını sarı yap
            tc_pr = cell._tc.get_or_add_tcPr()
            
            # Eski shading'i kaldır
            for shading in tc_pr.findall(qn('w:shd')):
                tc_pr.remove(shading)
            
            # Sarı arka plan ekle
            shading_elm = OxmlElement(qn('w:shd'))
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), 'FFFF00')  # Sarı
            
            tc_pr.append(shading_elm)
            print(f"      🟡 Hücre arka planı sarı yapıldı: {text}")
            return True
            
        except Exception as e:
            print(f"      ⚠ Sarı arka plan hatası: {e}")
            return False
    
    def format_tolerance_violation(self, cell, text: str, tolerance_results: List[bool] = None):
        """
        Tolerans dışı değerleri formatlar - önce sarı highlight, başarısız olursa kırmızı bold
        
        Args:
            cell: Word tablosu hücresi
            text: Yazılacak metin
            tolerance_results: Her değer için tolerans durumu listesi
        """
        try:
            print(f"      🔍 Tolerans dışı format uygulanıyor: {text}")
            
            # Önce sarı highlight dene (Stack Overflow yöntemi)
            if tolerance_results and len(tolerance_results) > 1:
                # Çoklu değer - sadece tolerans dışı olanları formatla
                out_of_tolerance_mask = [not result for result in tolerance_results]
                self.apply_yellow_highlight(cell, text, out_of_tolerance_mask)
                return True
            else:
                # Tek değer veya tüm değerler - hepsini formatla
                self.apply_yellow_highlight(cell, text, None)
                return True
                
        except Exception as e:
            print(f"      ⚠ Sarı highlight başarısız, kırmızı bold deneniyor: {e}")
            
            # Fallback 1: Kırmızı bold format
            try:
                if tolerance_results and len(tolerance_results) > 1:
                    out_of_tolerance_mask = [not result for result in tolerance_results]
                    self.apply_red_bold_format(cell, text, out_of_tolerance_mask)
                else:
                    self.apply_red_bold_format(cell, text, None)
                print(f"      🔴 Kırmızı bold başarılı")
                return True
            except Exception as e2:
                print(f"      ⚠ Kırmızı bold da başarısız, sarı arka plan deneniyor: {e2}")
                
                # Fallback 2: Sarı arka plan
                success = self.apply_yellow_highlight_fallback(cell, text)
                if success:
                    print(f"      🟡 Sarı arka plan başarılı")
                    return True
                else:
                    # Son çare: Sadece metni yaz
                    cell.text = text
                    print(f"      ⚠ [SON ÇARE] Sadece metin yazıldı: {text}")
                    return False

    def update_actual_values(self, karakterler: List[TeknikResimKarakteri]) -> bool:
        """Word tablosundaki ACTUAL kolonunu günceller (tolerans kontrolü ile)"""
        if not self.current_document:
            print("HATA: Önce Word dosyası yüklenmelidir")
            return False
        
        try:
            print("ACTUAL değerleri Word tablosuna yazılıyor (tolerans kontrolü ile)...")
            
            # Karakterleri item_no ile hızlı erişim için dict'e çevir
            karakter_dict = {k.item_no: k for k in karakterler}
            updated_count = 0
            tolerance_violations = 0
            
            for table_idx, table in enumerate(self.current_document.tables):
                print(f"  Tablo {table_idx + 1} kontrol ediliyor...")
                
                # Kolon indekslerini bul
                item_no_col_index, actual_col_index = self.find_table_columns(table)
                
                if actual_col_index is None:
                    print(f"    Tablo {table_idx + 1}'de ACTUAL kolonu bulunamadı")
                    continue
                
                if item_no_col_index is None:
                    print(f"    Tablo {table_idx + 1}'de ITEM NO kolonu bulunamadı")
                    continue
                
                print(f"    ITEM NO kolonu: {item_no_col_index}, ACTUAL kolonu: {actual_col_index}")
                
                # Veri satırlarını güncelle (header'ı atla)
                for row_idx, row in enumerate(table.rows[1:], start=1):
                    try:
                        # Güvenlik kontrolü - yeterli hücre var mı?
                        if len(row.cells) <= max(item_no_col_index, actual_col_index):
                            print(f"    ⚠ Satır {row_idx}: Yeterli hücre yok ({len(row.cells)} hücre)")
                            continue
                        
                        # ITEM NO'yu al
                        item_no_cell = row.cells[item_no_col_index]
                        item_no = item_no_cell.text.strip()
                        
                        # Bu ITEM NO'ya sahip karakter var mı ve actual değeri var mı?
                        if item_no in karakter_dict:
                            karakter = karakter_dict[item_no]
                            
                            if karakter.actual:
                                # ACTUAL hücresini güncelle
                                actual_cell = row.cells[actual_col_index]
                                actual_value = str(karakter.actual)
                                
                                # Tolerans kontrolü yap
                                tolerance_results, tolerance_status = self.check_multiple_values_tolerance(actual_value, karakter)
                                
                                if tolerance_results:
                                    # En az bir değer var
                                    any_out_of_tolerance = not all(tolerance_results)
                                    
                                    if any_out_of_tolerance:
                                        # Tolerans dışı değer(ler) var - özel format uygula
                                        success = self.format_tolerance_violation(actual_cell, actual_value, tolerance_results)
                                        
                                        tolerance_violations += 1
                                        format_type = "sarı highlight veya kırmızı bold"
                                        print(f"    🟡 {item_no}: {actual_value} (Tolerans dışı - {format_type})")
                                        print(f"      📊 {tolerance_status}")
                                    else:
                                        # Tüm değerler tolerans içinde - normal yaz
                                        actual_cell.text = actual_value
                                        print(f"    ✅ {item_no}: {actual_value} (Tolerans içinde)")
                                else:
                                    # Tolerans kontrolü yapılamadı - normal yaz
                                    actual_cell.text = actual_value
                                    print(f"    ○ {item_no}: {actual_value} (Tolerans kontrol edilemedi)")
                                
                                updated_count += 1
                            else:
                                print(f"    ○ {item_no}: Ölçüm değeri yok")
                        else:
                            # ITEM NO bulunamadı - sadece debug için yazdır
                            if item_no and item_no.upper().startswith('KN'):
                                print(f"    ? {item_no}: Karakterler listesinde bulunamadı")
                                
                    except Exception as e:
                        print(f"    ✗ Satır {row_idx} güncellenirken hata: {e}")
                        continue
            
            print(f"✓ Toplam {updated_count} ACTUAL değeri güncellendi")
            if tolerance_violations > 0:
                print(f"🟡 {tolerance_violations} tolerans dışı değer sarı highlight ile işaretlendi")
            
            return updated_count > 0
            
        except Exception as e:
            print(f"HATA: ACTUAL değerleri güncellenirken hata: {e}")
            return False
    
    def get_save_path(self, suggested_name: str = None) -> str:
        """Kullanıcıdan kaydetme yolunu alır"""
        if not suggested_name:
            original_name = os.path.basename(self.original_file_path) if self.original_file_path else "document.docx"
            name_without_ext = os.path.splitext(original_name)[0]
            suggested_name = f"{name_without_ext}_with_measurements.docx"
        
        try:
            # Modern tkinter için initialfile kullan
            save_path = filedialog.asksaveasfilename(
                title="Word Dosyasını Ölçüm Değerleriyle Kaydet",
                defaultextension=".docx",
                initialfile=suggested_name,
                filetypes=[
                    ("Word Dosyaları", "*.docx"),
                    ("Tüm Dosyalar", "*.*")
                ]
            )
        except Exception as e:
            print(f"initialfile hatası, fallback kullanılıyor: {e}")
            # Fallback - eski tkinter sürümleri için
            save_path = filedialog.asksaveasfilename(
                title=f"Word Dosyasını Kaydet (Önerilen: {suggested_name})",
                defaultextension=".docx",
                filetypes=[
                    ("Word Dosyaları", "*.docx"),
                    ("Tüm Dosyalar", "*.*")
                ]
            )
        
        return save_path
    
    def save_document(self, save_path: str) -> bool:
        """Dokümanı belirtilen yola kaydeder"""
        try:
            if not self.current_document:
                raise Exception("Kaydedilecek doküman yok")
            
            self.current_document.save(save_path)
            print(f"✓ Word dosyası kaydedildi: {save_path}")
            return True
            
        except Exception as e:
            print(f"HATA: Word dosyası kaydedilemedi: {e}")
            return False
    
    def save_with_actual_values(self, karakterler: List[TeknikResimKarakteri], save_path: str = None) -> str:
        """
        Tüm işlemi tek seferde yapar: ACTUAL değerlerini günceller ve kaydeder
        
        Args:
            karakterler: Ölçüm değerleri içeren karakter listesi
            save_path: Kaydetme yolu (None ise kullanıcıdan istenir)
            
        Returns:
            str: Kaydedilen dosyanın yolu
            
        Raises:
            Exception: Kaydetme işlemi başarısız olursa
        """
        try:
            print("Word dosyası ölçüm değerleriyle kaydediliyor (tolerans kontrolü ile)...")
            
            # 1. ACTUAL değerlerini güncelle (tolerans kontrolü ile)
            update_success = self.update_actual_values(karakterler)
            
            if not update_success:
                print("⚠ Hiçbir ACTUAL değeri güncellenmedi, yine de devam ediliyor...")
            
            # 2. Kaydetme yolunu belirle
            if not save_path:
                save_path = self.get_save_path()
                
                if not save_path:
                    raise Exception("Kaydetme işlemi iptal edildi")
            
            # 3. Dosyayı kaydet
            if not self.save_document(save_path):
                raise Exception("Dosya kaydetme işlemi başarısız")
            
            print(f"✓ Ölçüm değerleri Word dosyasına başarıyla aktarıldı: {save_path}")
            print("🟡 Tolerans dışı değerler sarı highlight ile işaretlendi")
            return save_path
            
        except Exception as e:
            error_msg = f"Word kaydetme işlemi başarısız: {str(e)}"
            print(f"HATA: {error_msg}")
            raise Exception(error_msg)
    
    def get_statistics(self, karakterler: List[TeknikResimKarakteri]) -> dict:
        """Ölçüm istatistiklerini döner (tolerans bilgileri dahil)"""
        total = len(karakterler)
        measured = len([k for k in karakterler if k.actual])
        unmeasured = total - measured
        
        # Tolerans istatistikleri
        tolerance_violations = 0
        tolerance_compliant = 0
        no_tolerance_defined = 0
        
        for karakter in karakterler:
            if karakter.actual:
                tolerance_results, _ = self.check_multiple_values_tolerance(karakter.actual, karakter)
                
                if tolerance_results:
                    if not all(tolerance_results):
                        tolerance_violations += 1
                    else:
                        tolerance_compliant += 1
                else:
                    no_tolerance_defined += 1
        
        return {
            'total': total,
            'measured': measured,
            'unmeasured': unmeasured,
            'completion_percentage': (measured / total * 100) if total > 0 else 0,
            'tolerance_violations': tolerance_violations,
            'tolerance_compliant': tolerance_compliant,
            'no_tolerance_defined': no_tolerance_defined
        }

# Test fonksiyonu
def test_word_save_as():
    """Word Save As servisini test eder (kırmızı bold format ile)"""
    from .word_reader import WordReaderService
    from .data_processor import DataProcessorService
    
    # Test dosyası
    test_file = "test_document.docx"  # Gerçek dosya yolunu buraya yazın
    
    if not os.path.exists(test_file):
        print(f"Test dosyası bulunamadı: {test_file}")
        return
    
    try:
        print("=== WORD SAVE AS TEST (Sarı Highlight ile) ===")
        
        # 1. Word dosyasını oku
        word_service = WordReaderService()
        df = DataProcessorService.from_word_tables(word_service, test_file)
        
        # 2. Karakterleri işle
        data_service = DataProcessorService()
        karakterler = data_service.process_dataframe(df)
        
        # 3. Test için çeşitli actual değerler ekle
        test_values = [
            "25.52",      # Normal değer
            "25.8",       # Tolerans dışı olabilir
            "25.4/25.6",  # Çoklu değer
            "25.2/25.9",  # Biri tolerans içi, biri dışı
        ]
        
        for i, karakter in enumerate(karakterler[:4]):
            if i < len(test_values):
                karakter.actual = test_values[i]
                print(f"Test actual eklendi: {karakter.item_no} = {karakter.actual}")
        
        # 4. İstatistikleri göster
        save_service = WordSaveAsService()
        stats = save_service.get_statistics(karakterler)
        print(f"\nİstatistikler:")
        print(f"  Toplam: {stats['total']}")
        print(f"  Ölçülen: {stats['measured']}")
        print(f"  Bekleyen: {stats['unmeasured']}")
        print(f"  Tamamlanma: %{stats['completion_percentage']:.1f}")
        print(f"  Tolerans dışı: {stats['tolerance_violations']}")
        print(f"  Tolerans içi: {stats['tolerance_compliant']}")
        
        # 5. Word Save As servisini test et
        if save_service.load_original_document(test_file):
            # Actual değerleriyle kaydet
            saved_path = save_service.save_with_actual_values(karakterler)
            print(f"\n✓ Test başarılı! Kaydedilen dosya: {saved_path}")
            print("🟡 Tolerans dışı değerler sarı highlight ile işaretlendi")
        else:
            print("Test başarısız: Dosya yüklenemedi")
            
    except Exception as e:
        print(f"Test hatası: {e}")

if __name__ == "__main__":
    test_word_save_as()