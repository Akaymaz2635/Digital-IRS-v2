# src/ui/components/document_viewer.py
"""
Word dokümanı görüntüleyici component'i - WebView ile HTML rendering
"""
import customtkinter as ctk
import os
import tempfile
import webbrowser
from pathlib import Path
from tkinter import messagebox

# Word to HTML conversion için
try:
    from docx import Document
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    print("mammoth kütüphanesi bulunamadı. pip install python-mammoth ile yükleyin")
    MAMMOTH_AVAILABLE = False

# WebView için
try:
    import tkinterweb
    WEBVIEW_AVAILABLE = True
except ImportError:
    print("tkinterweb bulunamadı. pip install tkinterweb ile yükleyin")
    WEBVIEW_AVAILABLE = False


class DocumentViewer(ctk.CTkFrame):
    """Word dokümanını görüntülemek için panel - WebView ile"""
    
    def __init__(self, parent):
        super().__init__(parent)
        
        self.current_html_file = None
        self.current_html_content = None
        self.current_zoom = 1.0
        
        self.setup_ui()
    
    def setup_ui(self):
        """Doküman görüntüleyici UI oluşturur"""
        # Başlık
        self._create_title()
        
        # Kontrol butonları
        self._create_control_buttons()
        
        # Ana görüntüleyici
        self._create_viewer()
    
    def _create_title(self):
        """Başlık bölümünü oluşturur"""
        title_label = ctk.CTkLabel(
            self,
            text="Word Dokümanı",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        title_label.pack(pady=(10, 5))
    
    def _create_control_buttons(self):
        """Kontrol butonlarını oluşturur"""
        button_frame = ctk.CTkFrame(self, fg_color="transparent")
        button_frame.pack(pady=5)
        
        # Doküman yükleme butonu
        self.load_button = ctk.CTkButton(
            button_frame,
            text="Dokümanı HTML'de Aç",
            command=self.open_in_browser,
            state="disabled",
            width=150
        )
        self.load_button.pack(side="left", padx=5)
        
        # Yenile butonu
        self.refresh_button = ctk.CTkButton(
            button_frame,
            text="Yenile",
            command=self.refresh_webview,
            state="disabled",
            width=80
        )
        self.refresh_button.pack(side="left", padx=5)
        
        # Zoom sıfırla butonu
        self.reset_zoom_button = ctk.CTkButton(
            button_frame,
            text="Zoom Reset",
            command=self.reset_zoom,
            state="disabled",
            width=100
        )
        self.reset_zoom_button.pack(side="left", padx=5)
    
    def _create_viewer(self):
        """Ana görüntüleyici alanını oluşturur"""
        if WEBVIEW_AVAILABLE:
            self._create_webview()
        else:
            self._create_fallback_textbox()
    
    def _create_webview(self):
        """WebView oluşturur"""
        # WebView frame
        self.web_frame = ctk.CTkFrame(self)
        self.web_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # WebView - scrollbar'lar ile
        self.webview = tkinterweb.HtmlFrame(
            self.web_frame,
            horizontal_scrollbar="auto",
            vertical_scrollbar="auto"
        )
        self.webview.pack(fill="both", expand=True)
        
        # Event bindings
        self.webview.bind("<Control-MouseWheel>", self.on_zoom)
        self.webview.focus_set()
        
        # Başlangıç HTML
        self._load_initial_html()
    
    def _create_fallback_textbox(self):
        """WebView yoksa fallback textbox oluşturur"""
        self.text_area = ctk.CTkTextbox(
            self,
            wrap="word",
            font=ctk.CTkFont(size=11)
        )
        self.text_area.pack(fill="both", expand=True, padx=10, pady=10)
        self.text_area.insert("1.0", "tkinterweb kütüphanesi bulunamadı.\nHTML görüntüleme için: pip install tkinterweb")
        self.text_area.configure(state="disabled")
    
    def _load_initial_html(self):
        """Başlangıç HTML'ini yükler"""
        if not WEBVIEW_AVAILABLE:
            return
            
        initial_html = """
        <html>
        <body style="font-family: Arial; padding: 20px; background-color: #2b2b2b; color: white;">
            <h3>Word Dokümanı Görüntüleyici</h3>
            <p>Word dosyası yüklendiğinde doküman içeriği burada görünecek.</p>
            <p><strong>Özellikler:</strong></p>
            <ul>
                <li>Tam HTML formatting</li>
                <li>Tablo yapısı korunur</li>
                <li>Scrollable içerik</li>
                <li>Ctrl + Mouse Wheel ile zoom</li>
            </ul>
        </body>
        </html>
        """
        self.webview.load_html(initial_html)
    
    def load_document(self, file_path: str):
        """Word dokümanını yükler ve görüntüler"""
        if not MAMMOTH_AVAILABLE:
            self._show_error("Mammoth kütüphanesi bulunamadı!")
            return
            
        try:
            print(f"Doküman WebView'da yükleniyor: {file_path}")
            
            if WEBVIEW_AVAILABLE:
                self._show_loading_message()
            
            # Word dokümanını işle
            html_content = self._convert_word_to_html(file_path)
            
            if html_content:
                # Styled HTML oluştur
                styled_html = self._create_styled_html(html_content, file_path)
                self.current_html_content = styled_html
                
                if WEBVIEW_AVAILABLE:
                    self.webview.load_html(styled_html)
                    self._enable_buttons()
                
                # HTML dosyası oluştur
                self._create_html_file(styled_html, file_path)
                
                print("✓ Doküman WebView'da başarıyla yüklendi")
            else:
                self._show_fallback_content(file_path)
                
        except Exception as e:
            error_msg = f"Doküman yükleme hatası: {str(e)}"
            print(error_msg)
            self._show_error(error_msg)
    
    def _show_loading_message(self):
        """Yükleniyor mesajını gösterir"""
        if WEBVIEW_AVAILABLE:
            loading_html = """
            <html>
            <body style="font-family: Arial; padding: 20px; background-color: #2b2b2b; color: white;">
                <h3>Yükleniyor...</h3>
                <p>Word dokümanı işleniyor, lütfen bekleyin...</p>
            </body>
            </html>
            """
            self.webview.load_html(loading_html)
    
    def _convert_word_to_html(self, file_path: str) -> str:
        """Word dosyasını HTML'e dönüştürür"""
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                
                if hasattr(result, 'value'):
                    return result.value
                elif hasattr(result, 'html'):
                    return result.html
                else:
                    raise Exception("HTML content bulunamadı")
                    
        except Exception as e:
            print(f"HTML dönüştürme hatası: {e}")
            return None
    
    def _create_styled_html(self, html_content: str, file_path: str) -> str:
        """HTML içeriğini stillendirme ile düzenler"""
        file_name = Path(file_path).stem
        
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Word Dokümanı - {file_name}</title>
            <style>
                body {{ 
                    font-family: 'Segoe UI', Arial, sans-serif; 
                    margin: 20px; 
                    line-height: 1.6;
                    background-color: #2b2b2b;
                    color: #ffffff;
                    overflow: auto;
                    min-width: 800px;
                }}
                h1, h2, h3 {{
                    color: #4fc3f7;
                    border-bottom: 2px solid #4fc3f7;
                    padding-bottom: 5px;
                }}
                table {{ 
                    border-collapse: collapse; 
                    width: 100%; 
                    min-width: 600px;
                    margin: 15px 0;
                    background-color: #3b3b3b;
                    border-radius: 5px;
                    overflow: hidden;
                }}
                th, td {{ 
                    border: 1px solid #555; 
                    padding: 12px 8px; 
                    text-align: left;
                }}
                th {{ 
                    background-color: #4fc3f7; 
                    font-weight: bold;
                    color: #000;
                }}
                tr:nth-child(even) {{
                    background-color: #404040;
                }}
                tr:hover {{
                    background-color: #505050;
                }}
                p {{
                    margin: 10px 0;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 30px;
                    padding: 20px;
                    background-color: #404040;
                    border-radius: 10px;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>📋 Word Dokümanı</h1>
                <h2>{file_name}</h2>
            </div>
            {html_content}
        </body>
        </html>
        """
        
        return styled_html
    
    def _show_fallback_content(self, file_path: str):
        """Fallback içerik gösterir"""
        if not WEBVIEW_AVAILABLE:
            return
            
        try:
            # Basit text extraction deneme
            doc = Document(file_path)
            
            # Text content topla
            text_content = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += f"<p>{para.text}</p>"
            
            # Tabloları HTML olarak ekle
            table_html = ""
            for i, table in enumerate(doc.tables):
                table_html += f"<h3>Tablo {i+1}</h3><table>"
                for row in table.rows:
                    table_html += "<tr>"
                    for cell in row.cells:
                        table_html += f"<td>{cell.text}</td>"
                    table_html += "</tr>"
                table_html += "</table>"
            
            fallback_html = f"""
            <html>
            <body style="font-family: Arial; padding: 20px; background-color: #2b2b2b; color: white;">
                <h2>Word Dokümanı (Text Modu)</h2>
                <p><em>HTML dönüştürme başarısız, basit text modu kullanılıyor.</em></p>
                {text_content}
                {table_html}
            </body>
            </html>
            """
            
            self.webview.load_html(fallback_html)
            print("Fallback text modu yüklendi")
            
        except Exception as e:
            print(f"Fallback content hatası: {e}")
            self._show_error("Doküman yüklenemedi")
    
    def _show_error(self, error_message: str):
        """Hata mesajını gösterir"""
        if WEBVIEW_AVAILABLE:
            error_html = f"""
            <html>
            <body style="font-family: Arial; padding: 20px; background-color: #2b2b2b; color: #ff6b6b;">
                <h3>Hata!</h3>
                <p>{error_message}</p>
                <p>Lütfen dosyayı kontrol edin ve tekrar deneyin.</p>
            </body>
            </html>
            """
            self.webview.load_html(error_html)
    
    def _enable_buttons(self):
        """Butonları aktif eder"""
        self.load_button.configure(state="normal")
        self.refresh_button.configure(state="normal")
        if WEBVIEW_AVAILABLE:
            self.reset_zoom_button.configure(state="normal")
    
    def _create_html_file(self, html_content: str, original_file: str):
        """HTML dosyası oluşturur"""
        try:
            temp_dir = tempfile.gettempdir()
            file_name = Path(original_file).stem
            html_file = os.path.join(temp_dir, f"{file_name}_preview.html")
            
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.current_html_file = html_file
            print(f"✓ HTML dosyası oluşturuldu: {html_file}")
            
        except Exception as e:
            print(f"HTML dosyası oluşturma hatası: {e}")
    
    def on_zoom(self, event):
        """Ctrl + Mouse Wheel ile zoom"""
        if not WEBVIEW_AVAILABLE:
            return
        
        try:
            delta = event.delta
            zoom_factor = 0.1
            
            if delta > 0:  # Zoom in
                self.current_zoom += zoom_factor
            else:  # Zoom out
                self.current_zoom -= zoom_factor
            
            # Zoom sınırları
            self.current_zoom = max(0.5, min(self.current_zoom, 3.0))
            
            # Zoom uygula
            self._apply_zoom()
            
        except Exception as e:
            print(f"Zoom hatası: {e}")
    
    def _apply_zoom(self):
        """Zoom seviyesini uygular"""
        if not WEBVIEW_AVAILABLE or not self.current_html_content:
            return
            
        try:
            # JavaScript ile zoom deneme
            zoom_script = f'document.body.style.zoom = "{self.current_zoom}";'
            self.webview.run_script(zoom_script)
            print(f"Zoom seviyesi: %{self.current_zoom*100:.0f}")
            
        except:
            # Fallback: HTML'i yeniden yükle
            self._apply_zoom_to_html()
    
    def _apply_zoom_to_html(self):
        """HTML'e zoom CSS'i ekleyerek yeniden yükle"""
        if not self.current_html_content or not WEBVIEW_AVAILABLE:
            return
            
        zoom_style = f"""
        <style>
        body {{ transform: scale({self.current_zoom}); transform-origin: top left; width: {100/self.current_zoom}%; }}
        </style>
        """
        
        # HTML'e zoom style'ı ekle
        zoomed_html = self.current_html_content.replace("</head>", f"{zoom_style}</head>")
        self.webview.load_html(zoomed_html)
    
    def reset_zoom(self):
        """Zoom'u sıfırla"""
        if WEBVIEW_AVAILABLE:
            self.current_zoom = 1.0
            self._apply_zoom()
            print("Zoom %100'e sıfırlandı")
    
    def refresh_webview(self):
        """WebView'ı yenile"""
        if WEBVIEW_AVAILABLE and self.current_html_content:
            self.webview.load_html(self.current_html_content)
            print("WebView yenilendi")
    
    def open_in_browser(self):
        """HTML dosyasını tarayıcıda aç"""
        if self.current_html_file and os.path.exists(self.current_html_file):
            try:
                webbrowser.open(f'file://{self.current_html_file}')
                print(f"HTML dosyası tarayıcıda açıldı: {self.current_html_file}")
            except Exception as e:
                messagebox.showerror("Hata", f"Tarayıcıda açılamadı: {str(e)}")
        else:
            messagebox.showwarning("Uyarı", "Önce bir doküman yükleyin!")
    
    def get_zoom_level(self) -> float:
        """Mevcut zoom seviyesini döner"""
        return self.current_zoom
    
    def set_zoom_level(self, zoom: float):
        """Zoom seviyesini ayarlar"""
        self.current_zoom = max(0.5, min(zoom, 3.0))
        self._apply_zoom()
    
    def is_document_loaded(self) -> bool:
        """Doküman yüklenip yüklenmediğini kontrol eder"""
        return self.current_html_content is not None